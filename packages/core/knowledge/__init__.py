"""Knowledge base management with Chroma vector DB."""

from __future__ import annotations

import hashlib
import json
import os
import shutil
import threading
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Callable
from urllib.parse import urlparse

import chromadb
from pydantic import BaseModel, Field

# Document processing imports
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Web scraping imports
try:
    import requests
    from bs4 import BeautifulSoup
    HAS_WEB_SCRAPING = True
except ImportError:
    HAS_WEB_SCRAPING = False


class WebSource(BaseModel):
    """A web source for knowledge ingestion."""

    id: str
    agent_id: str
    url: str
    name: str
    description: str = ""
    refresh_interval_hours: int = 24
    last_refreshed: str | None = None
    last_refresh_status: str = "pending"
    chunk_count: int = 0
    created_at: str = Field(default_factory=lambda: datetime.utcnow().isoformat())
    metadata: dict[str, Any] = Field(default_factory=dict)
    auto_refresh: bool = True
    selector: str | None = None  # CSS selector to extract specific content


class KnowledgeDocument(BaseModel):
    """A document in an agent's knowledge base."""

    id: str
    agent_id: str
    filename: str
    file_type: str
    file_size: int
    chunk_count: int
    uploaded_at: str = Field(default_factory=lambda: datetime.utcnow().isoformat())
    metadata: dict[str, Any] = Field(default_factory=dict)


class KnowledgeChunk(BaseModel):
    """A chunk of text from a document."""

    id: str
    document_id: str
    text: str
    chunk_index: int
    metadata: dict[str, Any] = Field(default_factory=dict)


class KnowledgeManager:
    """Manages knowledge bases for agents using Chroma."""

    def __init__(self, storage_path: str | None = None):
        if storage_path is None:
            storage_path = os.path.join(
                os.path.dirname(__file__), "..", "..", "..", "data", "knowledge"
            )
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)

        # Initialize Chroma client with persistent storage (new API)
        chroma_path = self.storage_path / "chroma"
        chroma_path.mkdir(parents=True, exist_ok=True)
        self._client = chromadb.PersistentClient(path=str(chroma_path))

        # Document metadata storage
        self._docs_path = self.storage_path / "documents.json"
        self._documents: dict[str, KnowledgeDocument] = {}
        self._load_documents()

        # Files storage
        self._files_path = self.storage_path / "files"
        self._files_path.mkdir(exist_ok=True)

    def _load_documents(self) -> None:
        """Load document metadata from storage."""
        import json

        if self._docs_path.exists():
            try:
                with open(self._docs_path) as f:
                    data = json.load(f)
                    for doc_data in data.get("documents", []):
                        doc = KnowledgeDocument(**doc_data)
                        self._documents[doc.id] = doc
            except Exception:
                self._documents = {}

    def _save_documents(self) -> None:
        """Save document metadata to storage."""
        import json

        data = {"documents": [doc.model_dump() for doc in self._documents.values()]}
        with open(self._docs_path, "w") as f:
            json.dump(data, f, indent=2)

    def _get_collection(self, agent_id: str) -> chromadb.Collection:
        """Get or create a Chroma collection for an agent."""
        collection_name = f"agent_{agent_id.replace('-', '_')}"
        return self._client.get_or_create_collection(
            name=collection_name,
            metadata={"agent_id": agent_id},
        )

    def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text from a document."""
        if file_type == "txt":
            return file_path.read_text(encoding="utf-8", errors="ignore")

        elif file_type == "pdf" and PdfReader:
            reader = PdfReader(str(file_path))
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)

        elif file_type == "docx" and DocxDocument:
            doc = DocxDocument(str(file_path))
            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)
            return "\n".join(text_parts)

        else:
            # Try to read as text
            try:
                return file_path.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                return ""

    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
        """Split text into overlapping chunks."""
        if not text:
            return []

        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]

            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind(".")
                last_newline = chunk.rfind("\n")
                break_point = max(last_period, last_newline)
                if break_point > chunk_size // 2:
                    chunk = chunk[: break_point + 1]
                    end = start + break_point + 1

            chunks.append(chunk.strip())
            start = end - overlap

        return [c for c in chunks if c]

    def add_document(
        self,
        agent_id: str,
        filename: str,
        content: bytes,
        metadata: dict[str, Any] | None = None,
    ) -> KnowledgeDocument:
        """Add a document to an agent's knowledge base."""
        # Determine file type
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

        # Generate document ID
        content_hash = hashlib.md5(content).hexdigest()[:8]
        doc_id = f"{agent_id}_{content_hash}"

        # Save file
        file_path = self._files_path / f"{doc_id}.{ext}"
        file_path.write_bytes(content)

        # Extract text
        text = self._extract_text(file_path, ext)

        # Chunk text
        chunks = self._chunk_text(text)

        # Get collection and add chunks
        collection = self._get_collection(agent_id)

        chunk_ids = []
        chunk_texts = []
        chunk_metadatas = []

        for i, chunk_text in enumerate(chunks):
            chunk_id = f"{doc_id}_chunk_{i}"
            chunk_ids.append(chunk_id)
            chunk_texts.append(chunk_text)
            chunk_metadatas.append({
                "document_id": doc_id,
                "chunk_index": i,
                "filename": filename,
                "agent_id": agent_id,
            })

        if chunk_ids:
            collection.add(
                ids=chunk_ids,
                documents=chunk_texts,
                metadatas=chunk_metadatas,
            )

        # Create document record
        doc = KnowledgeDocument(
            id=doc_id,
            agent_id=agent_id,
            filename=filename,
            file_type=ext,
            file_size=len(content),
            chunk_count=len(chunks),
            metadata=metadata or {},
        )
        self._documents[doc_id] = doc
        self._save_documents()

        return doc

    def list_documents(self, agent_id: str) -> list[KnowledgeDocument]:
        """List all documents for an agent."""
        return [doc for doc in self._documents.values() if doc.agent_id == agent_id]

    def get_document(self, document_id: str) -> KnowledgeDocument | None:
        """Get a document by ID."""
        return self._documents.get(document_id)

    def delete_document(self, document_id: str) -> bool:
        """Delete a document from the knowledge base."""
        doc = self._documents.get(document_id)
        if not doc:
            return False

        # Delete from Chroma
        collection = self._get_collection(doc.agent_id)
        try:
            # Get all chunk IDs for this document
            results = collection.get(where={"document_id": document_id})
            if results["ids"]:
                collection.delete(ids=results["ids"])
        except Exception:
            pass

        # Delete file
        for ext in ["txt", "pdf", "docx", "doc"]:
            file_path = self._files_path / f"{document_id}.{ext}"
            if file_path.exists():
                file_path.unlink()
                break

        # Remove from metadata
        del self._documents[document_id]
        self._save_documents()

        return True

    def query(
        self,
        agent_id: str,
        query_text: str,
        n_results: int = 5,
    ) -> list[dict[str, Any]]:
        """Query an agent's knowledge base."""
        collection = self._get_collection(agent_id)

        try:
            results = collection.query(
                query_texts=[query_text],
                n_results=n_results,
            )
        except Exception:
            return []

        # Format results
        formatted = []
        if results["documents"] and results["documents"][0]:
            for i, doc_text in enumerate(results["documents"][0]):
                metadata = results["metadatas"][0][i] if results["metadatas"] else {}
                distance = results["distances"][0][i] if results["distances"] else 0
                formatted.append({
                    "text": doc_text,
                    "metadata": metadata,
                    "relevance": 1 - distance,  # Convert distance to relevance
                })

        return formatted

    def clear_agent_knowledge(self, agent_id: str) -> int:
        """Clear all knowledge for an agent."""
        docs = self.list_documents(agent_id)
        count = 0
        for doc in docs:
            if self.delete_document(doc.id):
                count += 1
        return count

    # =========================================================================
    # Web Source Management
    # =========================================================================

    def _load_web_sources(self) -> None:
        """Load web sources from storage."""
        sources_path = self.storage_path / "web_sources.json"
        if sources_path.exists():
            try:
                with open(sources_path) as f:
                    data = json.load(f)
                    self._web_sources = {
                        s["id"]: WebSource(**s) for s in data.get("sources", [])
                    }
            except Exception:
                self._web_sources = {}
        else:
            self._web_sources = {}

    def _save_web_sources(self) -> None:
        """Save web sources to storage."""
        sources_path = self.storage_path / "web_sources.json"
        data = {"sources": [s.model_dump() for s in self._web_sources.values()]}
        with open(sources_path, "w") as f:
            json.dump(data, f, indent=2)

    def _fetch_url_content(self, url: str, selector: str | None = None) -> tuple[str, str]:
        """Fetch content from a URL.

        Returns: (text_content, title)
        """
        if not HAS_WEB_SCRAPING:
            raise RuntimeError("Web scraping not available. Install requests and beautifulsoup4.")

        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }

        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, "html.parser")

        # Get title
        title = soup.title.string if soup.title else urlparse(url).netloc

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        # Extract content
        if selector:
            elements = soup.select(selector)
            text = "\n\n".join(el.get_text(separator="\n", strip=True) for el in elements)
        else:
            # Try common content selectors
            main_content = (
                soup.find("main") or
                soup.find("article") or
                soup.find(class_="content") or
                soup.find(id="content") or
                soup.body
            )
            text = main_content.get_text(separator="\n", strip=True) if main_content else ""

        # Clean up text
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        text = "\n".join(lines)

        return text, title

    def add_web_source(
        self,
        agent_id: str,
        url: str,
        name: str | None = None,
        description: str = "",
        refresh_interval_hours: int = 24,
        selector: str | None = None,
        auto_refresh: bool = True,
    ) -> WebSource:
        """Add a web source and ingest its content."""
        if not HAS_WEB_SCRAPING:
            raise RuntimeError("Web scraping not available. Install requests and beautifulsoup4.")

        # Generate source ID
        url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
        source_id = f"web_{agent_id}_{url_hash}"

        # Check if already exists
        if source_id in self._web_sources:
            # Refresh existing source
            return self.refresh_web_source(source_id)

        # Fetch content
        try:
            text, title = self._fetch_url_content(url, selector)
            status = "success"
        except Exception as e:
            text = ""
            title = name or urlparse(url).netloc
            status = f"error: {str(e)[:100]}"

        # Use title as name if not provided
        if not name:
            name = title

        # Chunk and store content
        chunks = self._chunk_text(text)
        collection = self._get_collection(agent_id)

        chunk_ids = []
        chunk_texts = []
        chunk_metadatas = []

        for i, chunk_text in enumerate(chunks):
            chunk_id = f"{source_id}_chunk_{i}"
            chunk_ids.append(chunk_id)
            chunk_texts.append(chunk_text)
            chunk_metadatas.append({
                "source_id": source_id,
                "source_type": "web",
                "chunk_index": i,
                "url": url,
                "agent_id": agent_id,
            })

        if chunk_ids:
            collection.add(
                ids=chunk_ids,
                documents=chunk_texts,
                metadatas=chunk_metadatas,
            )

        # Create source record
        source = WebSource(
            id=source_id,
            agent_id=agent_id,
            url=url,
            name=name,
            description=description,
            refresh_interval_hours=refresh_interval_hours,
            last_refreshed=datetime.utcnow().isoformat(),
            last_refresh_status=status,
            chunk_count=len(chunks),
            auto_refresh=auto_refresh,
            selector=selector,
        )

        self._web_sources[source_id] = source
        self._save_web_sources()

        return source

    def refresh_web_source(self, source_id: str) -> WebSource:
        """Refresh a web source by re-fetching its content."""
        source = self._web_sources.get(source_id)
        if not source:
            raise ValueError(f"Web source '{source_id}' not found")

        # Delete old chunks
        collection = self._get_collection(source.agent_id)
        try:
            results = collection.get(where={"source_id": source_id})
            if results["ids"]:
                collection.delete(ids=results["ids"])
        except Exception:
            pass

        # Fetch new content
        try:
            text, _ = self._fetch_url_content(source.url, source.selector)
            status = "success"
        except Exception as e:
            text = ""
            status = f"error: {str(e)[:100]}"

        # Chunk and store
        chunks = self._chunk_text(text)

        chunk_ids = []
        chunk_texts = []
        chunk_metadatas = []

        for i, chunk_text in enumerate(chunks):
            chunk_id = f"{source_id}_chunk_{i}"
            chunk_ids.append(chunk_id)
            chunk_texts.append(chunk_text)
            chunk_metadatas.append({
                "source_id": source_id,
                "source_type": "web",
                "chunk_index": i,
                "url": source.url,
                "agent_id": source.agent_id,
            })

        if chunk_ids:
            collection.add(
                ids=chunk_ids,
                documents=chunk_texts,
                metadatas=chunk_metadatas,
            )

        # Update source record
        source.last_refreshed = datetime.utcnow().isoformat()
        source.last_refresh_status = status
        source.chunk_count = len(chunks)
        self._save_web_sources()

        return source

    def list_web_sources(self, agent_id: str | None = None) -> list[WebSource]:
        """List web sources, optionally filtered by agent."""
        if agent_id:
            return [s for s in self._web_sources.values() if s.agent_id == agent_id]
        return list(self._web_sources.values())

    def get_web_source(self, source_id: str) -> WebSource | None:
        """Get a web source by ID."""
        return self._web_sources.get(source_id)

    def delete_web_source(self, source_id: str) -> bool:
        """Delete a web source and its chunks."""
        source = self._web_sources.get(source_id)
        if not source:
            return False

        # Delete chunks from Chroma
        collection = self._get_collection(source.agent_id)
        try:
            results = collection.get(where={"source_id": source_id})
            if results["ids"]:
                collection.delete(ids=results["ids"])
        except Exception:
            pass

        # Remove from storage
        del self._web_sources[source_id]
        self._save_web_sources()

        return True

    def get_sources_needing_refresh(self) -> list[WebSource]:
        """Get web sources that need to be refreshed."""
        now = datetime.utcnow()
        sources_to_refresh = []

        for source in self._web_sources.values():
            if not source.auto_refresh:
                continue

            if not source.last_refreshed:
                sources_to_refresh.append(source)
                continue

            last_refresh = datetime.fromisoformat(source.last_refreshed)
            next_refresh = last_refresh + timedelta(hours=source.refresh_interval_hours)

            if now >= next_refresh:
                sources_to_refresh.append(source)

        return sources_to_refresh

    def refresh_all_due_sources(self) -> dict[str, str]:
        """Refresh all sources that are due for refresh.

        Returns: Dict mapping source_id to status
        """
        results = {}
        for source in self.get_sources_needing_refresh():
            try:
                self.refresh_web_source(source.id)
                results[source.id] = "success"
            except Exception as e:
                results[source.id] = f"error: {str(e)[:100]}"
        return results


class KnowledgeScheduler:
    """Background scheduler for refreshing web sources."""

    def __init__(self, manager: KnowledgeManager, check_interval_seconds: int = 3600):
        self.manager = manager
        self.check_interval = check_interval_seconds
        self._running = False
        self._thread: threading.Thread | None = None
        self._callbacks: list[Callable[[str, str], None]] = []

    def add_callback(self, callback: Callable[[str, str], None]) -> None:
        """Add a callback to be called when a source is refreshed.

        Callback signature: (source_id, status)
        """
        self._callbacks.append(callback)

    def _run_loop(self) -> None:
        """Main scheduler loop."""
        while self._running:
            try:
                results = self.manager.refresh_all_due_sources()
                for source_id, status in results.items():
                    for callback in self._callbacks:
                        try:
                            callback(source_id, status)
                        except Exception:
                            pass
            except Exception:
                pass

            # Sleep in small increments to allow quick shutdown
            for _ in range(self.check_interval):
                if not self._running:
                    break
                time.sleep(1)

    def start(self) -> None:
        """Start the scheduler in a background thread."""
        if self._running:
            return

        self._running = True
        self._thread = threading.Thread(target=self._run_loop, daemon=True)
        self._thread.start()

    def stop(self) -> None:
        """Stop the scheduler."""
        self._running = False
        if self._thread:
            self._thread.join(timeout=5)
            self._thread = None

    def is_running(self) -> bool:
        """Check if the scheduler is running."""
        return self._running


# Singleton instances
_knowledge_manager: KnowledgeManager | None = None
_knowledge_scheduler: KnowledgeScheduler | None = None


def get_knowledge_manager() -> KnowledgeManager:
    """Get the knowledge manager singleton."""
    global _knowledge_manager
    if _knowledge_manager is None:
        _knowledge_manager = KnowledgeManager()
        # Load web sources
        _knowledge_manager._web_sources = {}
        _knowledge_manager._load_web_sources()
    return _knowledge_manager


def get_knowledge_scheduler() -> KnowledgeScheduler:
    """Get the knowledge scheduler singleton."""
    global _knowledge_scheduler
    if _knowledge_scheduler is None:
        _knowledge_scheduler = KnowledgeScheduler(get_knowledge_manager())
    return _knowledge_scheduler


def start_knowledge_scheduler() -> None:
    """Start the background knowledge scheduler."""
    scheduler = get_knowledge_scheduler()
    if not scheduler.is_running():
        scheduler.start()


def stop_knowledge_scheduler() -> None:
    """Stop the background knowledge scheduler."""
    global _knowledge_scheduler
    if _knowledge_scheduler:
        _knowledge_scheduler.stop()


__all__ = [
    "KnowledgeDocument",
    "WebSource",
    "KnowledgeManager",
    "KnowledgeScheduler",
    "get_knowledge_manager",
    "get_knowledge_scheduler",
    "start_knowledge_scheduler",
    "stop_knowledge_scheduler",
]
